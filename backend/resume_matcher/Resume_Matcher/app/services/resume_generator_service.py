import io
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches


class ResumeGeneratorService:
    """Generate a downloadable Word document from enhanced resume data."""

    def generate_resume_docx(self, enhanced_resume: Dict) -> io.BytesIO:
        doc = Document()

        # Default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # PERSONAL INFO
        self._add_personal_info(doc, enhanced_resume.get("personal_info", {}))

        # OPTIONAL SUMMARY
        if enhanced_resume.get("summary"):
            self._add_section(doc, "Professional Summary", enhanced_resume["summary"])

        # EXPERIENCE
        if enhanced_resume.get("experience"):
            self._add_experience_section(doc, enhanced_resume["experience"])

        # EDUCATION
        if enhanced_resume.get("education"):
            self._add_education_section(doc, enhanced_resume["education"])

        # SKILLS
        if enhanced_resume.get("skills"):
            self._add_skills_section(doc, enhanced_resume["skills"])

        # ACHIEVEMENTS
        if enhanced_resume.get("achievements"):
            self._add_achievements_section(doc, enhanced_resume["achievements"])

        # AI ENHANCEMENT SUGGESTIONS
        if enhanced_resume.get("enhancements"):
            self._add_enhancements_section(doc, enhanced_resume["enhancements"])

        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # PERSONAL INFO
    def _add_personal_info(self, doc: Document, info: Dict):
        name = info.get("name", "Your Name")

        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(16)

        # Contact info
        contact = []
        for key in ["email", "phone", "location"]:
            if info.get(key):
                contact.append(info[key])

        if contact:
            p = doc.add_paragraph(" | ".join(contact))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].font.size = Pt(10)

        doc.add_paragraph("_" * 50)  # separator
        doc.add_paragraph()

    # GENERIC SECTION
    def _add_section(self, doc: Document, title: str, content: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(content)
        doc.add_paragraph()

    # EXPERIENCE
    def _add_experience_section(self, doc: Document, experiences: List[Dict]):
        p = doc.add_paragraph()
        run = p.add_run("Professional Experience")
        run.bold = True
        run.font.size = Pt(12)

        for exp in experiences:
            p = doc.add_paragraph()
            title = exp.get("title", "Title")
            company = exp.get("company", "Company")
            p.add_run(f"{title} at {company}").bold = True

            # Duration
            start = exp.get("start_date")
            end = exp.get("end_date")
            if start or end:
                dates = f"{start or ''} - {end or ''}"
                p = doc.add_paragraph()
                p.add_run(dates).italic = True

            # Responsibilities
            responsibilities = exp.get("responsibilities", [])
            for r in responsibilities:
                doc.add_paragraph(r, style="List Bullet")

            doc.add_paragraph()

    # EDUCATION
    def _add_education_section(self, doc: Document, education: List[Dict]):
        p = doc.add_paragraph()
        run = p.add_run("Education")
        run.bold = True
        run.font.size = Pt(12)

        for edu in education:
            degree = edu.get("degree", "Degree")
            institution = edu.get("institution", "Institution")

            p = doc.add_paragraph()
            p.add_run(f"{degree} — {institution}").bold = True

            if edu.get("year_of_completion"):
                year = edu["year_of_completion"]
                doc.add_paragraph(str(year)).italic = True

            doc.add_paragraph()

    # SKILLS
    def _add_skills_section(self, doc: Document, skills: List[str]):
        p = doc.add_paragraph()
        run = p.add_run("Skills")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(", ".join(skills))
        doc.add_paragraph()

    # ACHIEVEMENTS
    def _add_achievements_section(self, doc: Document, achievements: List[str]):
        p = doc.add_paragraph()
        run = p.add_run("Achievements")
        run.bold = True
        run.font.size = Pt(12)

        for ach in achievements:
            doc.add_paragraph(ach, style="List Bullet")

        doc.add_paragraph()

    # ENHANCEMENTS SECTION
    def _add_enhancements_section(self, doc: Document, enhancements: Dict):
        p = doc.add_paragraph()
        run = p.add_run("AI-Powered Enhancement Suggestions")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        for key, value in enhancements.items():
            formatted_key = key.replace("_", " ").title()

            p = doc.add_paragraph()
            p.add_run(f"{formatted_key}:").bold = True

            if isinstance(value, list):
                for item in value:
                    doc.add_paragraph(item, style="List Bullet")
            elif isinstance(value, str):
                doc.add_paragraph(value)

        doc.add_paragraph()
        doc.add_paragraph(
            "Note: These suggestions are AI-generated. Modify or remove as needed before using your resume."
        ).runs[0].italic = True