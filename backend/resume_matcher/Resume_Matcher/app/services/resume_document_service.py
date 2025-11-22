import io 
from typing import Dict, List 

from docx import Document 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

class ResumeDocumentServivce:
  """Creates a downloadable DOCX resume from enhanced JSON data."""

  def generate_docx(self, enhanced_resume: Dict) -> io.BytesIO:
    doc = Document()

    # Set default font 
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Personal Info 
    self._add_personal_info(doc, enhanced_resume.get("personal_info", {}))

    # Summary 
    if enhanced_resume.get("summary"):
      self._add_section(doc, "Professional Summary",enhanced_resume["summary"])

    # Experience 
    if enhanced_resume.get("experience"):
      self._add_experience_section(doc, enhanced_resume["experience"])

    # Education 
    if enhanced_resume.get("education"):
      self._add_education_section(doc, enhanced_resume['education'])

    # Skills 
    if enhanced_resume.get("skills"):
      self._add_skills_section(doc, enhanced_resume["skills"])

    # Achievements 
    if enhanced_resume.get("achievements"):
      self._add_achievements_section(doc, enhanced_resume["achievements"])

    # Enhancements 
    if enhanced_resume.get("enhancements"):
      self._add_enhancements_section(doc, enhanced_resume["enhancements"])

    # Return DOCX buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
  
  # SECTIONS 
  def _add_personal_info(self, doc, info:Dict):
    name = info.get("name", "Your Name")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(16)

    contact = []
    for field in ['email', 'phone','location']:
      if info.get(field):
        contact.append(info[field])

    if contact:
      p = doc.add_paragraph(" | ".join(contact))
      p.alignment = WD_ALIGN_PARAGRAPH.CENTER
      p.run[0].font.size = Pt(10)    

    doc.add_paragraph("_"*55)
    doc.add_paragraph()

  def _add_section(self, doc, title: str, content: str):
    run = doc.add_paragraph().add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(content)
    doc.add_paragraph()

  def _add_experience_section(self, doc, experiences: List[Dict]):
    run = doc.add_paragraph().add_run("Professional Experience")
    run.bold = True
    run.font.size = Pt(12)

    for exp in experiences:
      # Title + Company 
      p = doc.add_paragraph()
      p.add_run(f"{exp.get('title', '')} at {exp.get('company', '')}").bold = True

      # Date 
      start = exp.get("start_date")
      end = exp.get("end_date")
      if start or end:
        p = doc.add_paragraph()
        p.add_run(f"{start or ''} - {end or ''}").italic = True

      # Responsibilities
      for r in exp.get("responsibilities", []):
        doc.add_paragraph(r, style = "List Bullet")

      doc.add_paragraph()

  def _add_education_section(self, doc, educaiton: List[Dict]):
    run = doc.add_paragraph().add_run("Education")
    run.bold = True
    run.font.size = Pt(12)

    for edu in educaiton:
      degree = edu.get("degree","Degree")
      institution = edu.get("institution","Institution")

      p = doc.add_paragraph()
      p.add_run(f"{degree}- {institution}").bold = True

      if edu.get("year_of_completion"):
        doc.add_paragraph(str(edu["year_of_completion"])).itlaic = True

      doc.add_paragraph()

  def _add_skills_section(self, doc, skills: List[str]):
    run = doc.add_paragraph().add_run("Skills")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(", ".join(skills))
    doc.add_paragraph()

  def _add_achievements_section(self, doc, achievements: List[str]):
    run = doc.add_paragraph().add_run("Achievements")
    run.bold = True
    run.font.size = Pt(12)

    for ach in achievements:
      doc.add_paragraph(ach, style = "List Bullet")

    doc.add_paragraph()

  def _add_enhancements_section(self, doc, enhancements: Dict):
    p = doc.add_paragraph()
    run = p.add_run("AI Enhancement Suggenstions")
    run.bold = True 
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0,102,255)

    for key, value in enhancements.items():
      key = key.replace("_"," ").title()

      p = doc.add_paragraph()
      p.add_run(f"{key}:").bold = True

      if isinstance(value, list):
        for item in value:
          doc.add_paragraph(item, style = "List Bullet")
      else:
        doc.add_paragraph(str(value))
    
    doc.add_paragraph()